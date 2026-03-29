#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import sys
import platform
import subprocess
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Please run: pip install python-docx")
    sys.exit(1)

try:
    import PyPDF2
except ImportError:
    print("Warning: PyPDF2 not installed. PDF extraction may fail. Run: pip install PyPDF2")

def extract_docx(file_path):
    """Extracts text from a .docx file."""
    try:
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_text))
        return "\n".join(content)
    except Exception as e:
        return f"Error extracting {os.path.basename(file_path)}: {e}"

def extract_pdf_pypdf2(file_path):
    """Extracts text from a .pdf file using PyPDF2 (cross-platform)."""
    try:
        import PyPDF2
        content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    content.append(text)
        return "\n".join(content)
    except Exception as e:
        return f"Error extracting {os.path.basename(file_path)}: {e}"

def extract_pdf_pdftotext(file_path, temp_dir):
    """Extracts text from a .pdf file using pdftotext (Linux/Mac)."""
    try:
        output_path = os.path.join(temp_dir, os.path.basename(file_path) + ".txt")
        subprocess.run(["pdftotext", "-layout", file_path, output_path], check=True, capture_output=True, text=True)
        with open(output_path, 'r', encoding='utf-8') as f:
            return f.read()
    except subprocess.CalledProcessError as e:
        return f"Error extracting {os.path.basename(file_path)}: {e.stderr}"
    except Exception as e:
        return f"Error processing {os.path.basename(file_path)}: {e}"

def extract_pdf(file_path, temp_dir):
    """Extracts text from a .pdf file using available tools."""
    # Try PyPDF2 first (cross-platform)
    try:
        import PyPDF2
        return extract_pdf_pypdf2(file_path)
    except ImportError:
        pass
    
    # Fall back to pdftotext (Linux/Mac)
    if platform.system() != "Windows":
        return extract_pdf_pdftotext(file_path, temp_dir)
    
    return f"Error: Unable to extract {os.path.basename(file_path)}. Please install PyPDF2: pip install PyPDF2"

def extract_doc_libreoffice(file_path, temp_dir):
    """Extracts text from a .doc file using LibreOffice (Linux/Mac)."""
    try:
        subprocess.run(["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, file_path], 
                      check=True, capture_output=True, text=True)
        output_path = os.path.join(temp_dir, os.path.splitext(os.path.basename(file_path))[0] + ".txt")
        if not os.path.exists(output_path):
            return f"Error: Conversion failed for {os.path.basename(file_path)}. Output file not found."
        with open(output_path, 'r', encoding='utf-8') as f:
            return f.read()
    except subprocess.CalledProcessError as e:
        return f"Error extracting {os.path.basename(file_path)}: {e.stderr}"
    except Exception as e:
        return f"Error processing {os.path.basename(file_path)}: {e}"

def extract_doc_python_docx(file_path):
    """Attempts to extract text from .doc files using python-docx (may have limited support)."""
    try:
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        return "\n".join(content)
    except Exception as e:
        return f"Error extracting {os.path.basename(file_path)}: {e}"

def extract_doc(file_path, temp_dir):
    """Extracts text from a .doc file using available tools."""
    # Try python-docx first (cross-platform, but limited .doc support)
    result = extract_doc_python_docx(file_path)
    if not result.startswith("Error"):
        return result
    
    # Fall back to LibreOffice (Linux/Mac)
    if platform.system() != "Windows":
        return extract_doc_libreoffice(file_path, temp_dir)
    
    return f"Warning: Limited .doc support on Windows. Please convert to .docx format or install LibreOffice."

def main():
    parser = argparse.ArgumentParser(description="Build a knowledge base from a directory of documents.")
    parser.add_argument("source_dir", help="The source directory containing .doc, .docx, and .pdf files.")
    parser.add_argument("output_file", help="The path for the final consolidated knowledge base file.")
    args = parser.parse_args()

    if not os.path.isdir(args.source_dir):
        print(f"Error: Source directory not found at {args.source_dir}")
        return

    temp_dir = os.path.join(os.path.expanduser("~"), ".kb_builder_temp")
    os.makedirs(temp_dir, exist_ok=True)

    all_content = ""
    file_count = 0
    error_count = 0

    print(f"Starting knowledge base build from: {args.source_dir}")
    print(f"Platform: {platform.system()} {platform.release()}")

    for filename in sorted(os.listdir(args.source_dir)):
        file_path = os.path.join(args.source_dir, filename)
        if not os.path.isfile(file_path):
            continue
            
        content = ""
        if filename.endswith(".docx"):
            print(f"- Processing .docx: {filename}")
            content = extract_docx(file_path)
        elif filename.endswith(".pdf"):
            print(f"- Processing .pdf: {filename}")
            content = extract_pdf(file_path, temp_dir)
        elif filename.endswith(".doc"):
            print(f"- Processing .doc: {filename}")
            content = extract_doc(file_path, temp_dir)
        
        if content:
            if content.startswith("Error") or content.startswith("Warning"):
                print(f"  {content}")
                error_count += 1
            else:
                all_content += f"\n\n--- START OF {filename} ---\n\n{content}\n\n--- END OF {filename} ---\n\n"
                file_count += 1

    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(args.output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    with open(args.output_file, 'w', encoding='utf-8') as f:
        f.write(f"# Knowledge Base\n\n")
        f.write(f"This knowledge base was automatically generated from {file_count} documents.\n")
        if error_count > 0:
            f.write(f"Note: {error_count} files could not be processed.\n")
        f.write(all_content)

    # Cleanup temp directory
    try:
        import shutil
        shutil.rmtree(temp_dir)
    except:
        pass

    print(f"\n✓ Knowledge base built successfully!")
    print(f"  - Documents processed: {file_count}")
    if error_count > 0:
        print(f"  - Documents failed: {error_count}")
    print(f"  - Output file: {args.output_file}")
    if os.path.exists(args.output_file):
        print(f"  - Size: {os.path.getsize(args.output_file) / 1024 / 1024:.2f} MB")

if __name__ == "__main__":
    main()
