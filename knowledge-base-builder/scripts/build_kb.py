#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import subprocess
import argparse
from docx import Document

def extract_docx(file_path):
    """Extracts text from a .docx file."""
    try:
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_text))
        return "\n".join(content)
    except Exception as e:
        return f"Error extracting {os.path.basename(file_path)}: {e}"

def extract_pdf(file_path, temp_dir):
    """Extracts text from a .pdf file using pdftotext."""
    try:
        output_path = os.path.join(temp_dir, os.path.basename(file_path) + ".txt")
        subprocess.run(["pdftotext", "-layout", file_path, output_path], check=True, capture_output=True, text=True)
        with open(output_path, 'r', encoding='utf-8') as f:
            return f.read()
    except subprocess.CalledProcessError as e:
        return f"Error extracting {os.path.basename(file_path)}: {e.stderr}"
    except Exception as e:
        return f"Error processing {os.path.basename(file_path)}: {e}"

def extract_doc(file_path, temp_dir):
    """Extracts text from a .doc file by converting to txt with LibreOffice."""
    try:
        subprocess.run(["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, file_path], check=True, capture_output=True, text=True)
        output_path = os.path.join(temp_dir, os.path.splitext(os.path.basename(file_path))[0] + ".txt")
        if not os.path.exists(output_path):
             return f"Error: Conversion failed for {os.path.basename(file_path)}. Output file not found."
        with open(output_path, 'r', encoding='utf-8') as f:
            return f.read()
    except subprocess.CalledProcessError as e:
        return f"Error extracting {os.path.basename(file_path)}: {e.stderr}"
    except Exception as e:
        return f"Error processing {os.path.basename(file_path)}: {e}"

def main():
    parser = argparse.ArgumentParser(description="Build a knowledge base from a directory of documents.")
    parser.add_argument("source_dir", help="The source directory containing .doc, .docx, and .pdf files.")
    parser.add_argument("output_file", help="The path for the final consolidated knowledge base file.")
    args = parser.parse_args()

    if not os.path.isdir(args.source_dir):
        print(f"Error: Source directory not found at {args.source_dir}")
        return

    temp_dir = "/tmp/kb_builder_temp"
    os.makedirs(temp_dir, exist_ok=True)

    all_content = ""
    file_count = 0

    print(f"Starting knowledge base build from: {args.source_dir}")

    for filename in sorted(os.listdir(args.source_dir)):
        file_path = os.path.join(args.source_dir, filename)
        content = ""
        if filename.endswith(".docx"):
            print(f"- Processing .docx: {filename}")
            content = extract_docx(file_path)
        elif filename.endswith(".pdf"):
            print(f"- Processing .pdf: {filename}")
            content = extract_pdf(file_path, temp_dir)
        elif filename.endswith(".doc"):
            print(f"- Processing .doc: {filename}")
            content = extract_doc(file_path, temp_dir)
        
        if content:
            all_content += f"\n\n--- START OF {filename} ---\n\n{content}\n\n--- END OF {filename} ---\n\n"
            file_count += 1

    with open(args.output_file, 'w', encoding='utf-8') as f:
        f.write(f"# Knowledge Base\n\n")
        f.write(f"This knowledge base was automatically generated from {file_count} documents found in the source directory.\n")
        f.write(all_content)

    # Cleanup temp directory
    subprocess.run(["rm", "-rf", temp_dir])

    print(f"\n\u2713 Knowledge base built successfully!")
    print(f"  - Documents processed: {file_count}")
    print(f"  - Output file: {args.output_file}")
    print(f"  - Size: {os.path.getsize(args.output_file) / 1024 / 1024:.2f} MB")

if __name__ == "__main__":
    main()
